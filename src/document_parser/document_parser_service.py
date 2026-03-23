import os
import logging
import tempfile
from typing import Union, Dict, Any, List
from pathlib import Path
import json

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF未安装，PDF解析功能将不可用。请安装: pip install pymupdf")

try:
    import pandas as pd
    import openpyxl
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logging.warning("pandas或openpyxl未安装，Excel解析功能将不可用。请安装: pip install pandas openpyxl")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx未安装，Word解析功能将不可用。请安装: pip install python-docx")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParserService:
    def __init__(self):
        self.ocr_service = None
    
    def _get_ocr_service(self):
        if self.ocr_service is None:
            from ocr.ocr_service import OCRService
            self.ocr_service = OCRService()
        return self.ocr_service
    
    def parse_pdf(
        self, 
        file_path: Union[str, Path],
        extract_images: bool = False,
        enable_ocr: bool = False,
        force_ocr: bool = False
    ) -> Dict[str, Any]:
        if not PYMUPDF_AVAILABLE:
            raise ImportError("请安装PyMuPDF: pip install pymupdf")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF文件不存在: {file_path}")
        
        logger.info(f"开始解析PDF: {file_path}")
        
        doc = fitz.open(file_path)
        
        metadata = {
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "creator": doc.metadata.get("creator", ""),
            "producer": doc.metadata.get("producer", ""),
            "creationDate": doc.metadata.get("creationDate", ""),
            "modDate": doc.metadata.get("modDate", ""),
            "total_pages": len(doc)
        }
        
        toc = doc.get_toc()
        
        pages_content = []
        all_text = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            text = page.get_text("text")
            
            blocks = page.get_text("dict")["blocks"]
            structured_blocks = []
            
            for block in blocks:
                if "lines" in block:
                    block_text = "\n".join([
                        " ".join([span["text"] for span in line["spans"]])
                        for line in block["lines"]
                    ])
                    structured_blocks.append({
                        "type": "text",
                        "bbox": block["bbox"],
                        "text": block_text
                    })
            
            page_data = {
                "page_number": page_num + 1,
                "text": text,
                "blocks": structured_blocks,
                "method": "direct_extraction"
            }
            
            if extract_images:
                images = self._extract_images_from_page(page, page_num)
                page_data["images"] = images
            
            if enable_ocr and not text.strip():
                logger.info(f"第 {page_num + 1} 页无文本，尝试OCR...")
                try:
                    ocr_service = self._get_ocr_service()
                    ocr_result = ocr_service.extract_text_from_pdf_smart(
                        file_path, 
                        force_ocr=True
                    )
                    if ocr_result.get("success") and ocr_result.get("pages"):
                        page_ocr_data = ocr_result["pages"][page_num]
                        page_data["ocr_text"] = page_ocr_data.get("text", "")
                        page_data["method"] = "ocr"
                except Exception as e:
                    logger.warning(f"第 {page_num + 1} 页OCR失败: {str(e)}")
            
            pages_content.append(page_data)
            all_text.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
        
        doc.close()
        
        return {
            "success": True,
            "file_type": "pdf",
            "metadata": metadata,
            "toc": toc,
            "pages": pages_content,
            "full_text": "\n\n".join(all_text),
            "total_pages": len(pages_content)
        }
    
    def _extract_images_from_page(self, page, page_num: int) -> List[Dict[str, Any]]:
        images = []
        image_list = page.get_images()
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = page.parent.extract_image(xref)
            
            images.append({
                "page": page_num + 1,
                "index": img_index,
                "width": base_image["width"],
                "height": base_image["height"],
                "colorspace": base_image["colorspace"],
                "image_data": base_image["image"]
            })
        
        return images
    
    def parse_excel(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        if not PANDAS_AVAILABLE:
            raise ImportError("请安装pandas和openpyxl: pip install pandas openpyxl")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Excel文件不存在: {file_path}")
        
        logger.info(f"开始解析Excel: {file_path}")
        
        excel_file = pd.ExcelFile(file_path)
        
        sheets_data = []
        all_text = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            sheet_dict = {
                "sheet_name": sheet_name,
                "shape": df.shape,
                "columns": df.columns.tolist(),
                "data": df.to_dict('records'),
                "text_representation": df.to_string()
            }
            
            sheets_data.append(sheet_dict)
            all_text.append(f"=== 工作表: {sheet_name} ===\n{df.to_string()}")
        
        return {
            "success": True,
            "file_type": "excel",
            "sheets": sheets_data,
            "full_text": "\n\n".join(all_text),
            "total_sheets": len(sheets_data)
        }
    
    def parse_word(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("请安装python-docx: pip install python-docx")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Word文件不存在: {file_path}")
        
        logger.info(f"开始解析Word文档: {file_path}")
        
        doc = Document(file_path)
        
        paragraphs = []
        all_text = []
        
        for i, para in enumerate(doc.paragraphs):
            para_data = {
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else None
            }
            paragraphs.append(para_data)
            all_text.append(para.text)
        
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            tables.append({
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data
            })
        
        return {
            "success": True,
            "file_type": "word",
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": "\n".join(all_text),
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables)
        }
    
    def parse_document(
        self, 
        file_path: Union[str, Path],
        **kwargs
    ) -> Dict[str, Any]:
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self.parse_pdf(
                file_path,
                extract_images=kwargs.get('extract_images', False),
                enable_ocr=kwargs.get('enable_ocr', False),
                force_ocr=kwargs.get('force_ocr', False)
            )
        elif suffix in ['.xlsx', '.xls']:
            return self.parse_excel(file_path)
        elif suffix in ['.docx', '.doc']:
            return self.parse_word(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")
    
    def extract_text_from_document(
        self, 
        file_path: Union[str, Path],
        **kwargs
    ) -> str:
        result = self.parse_document(file_path, **kwargs)
        
        if result.get("success"):
            return result.get("full_text", "")
        else:
            raise Exception(f"文档解析失败: {result.get('error', '未知错误')}")

if __name__ == "__main__":
    parser = DocumentParserService()
    
    test_files = {
        "pdf": "test.pdf",
        "excel": "test.xlsx",
        "word": "test.docx"
    }
    
    for file_type, file_path in test_files.items():
        if os.path.exists(file_path):
            print(f"\n=== 解析 {file_type.upper()} 文件 ===")
            result = parser.parse_document(file_path)
            print(f"成功: {result['success']}")
            print(f"文本长度: {len(result.get('full_text', ''))}")
